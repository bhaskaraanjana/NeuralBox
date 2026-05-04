from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'doc' / 'NeuralBox_Builder_Club_Presentation.docx'
ASSET_DIR = ROOT / 'doc' / '.presentation_assets'
ASSET_DIR.mkdir(parents=True, exist_ok=True)

COLORS = {
    'navy': '07111f',
    'ink': '0f172a',
    'blue': '2563eb',
    'cyan': '38bdf8',
    'green': '34d399',
    'amber': 'fbbf24',
    'red': 'f87171',
    'muted': '64748b',
    'paper': 'f8fafc',
    'line': 'cbd5e1',
    'white': 'ffffff',
}

def rgb(hex_color):
    hex_color = hex_color.strip('#')
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, color='0f172a', size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = rgb(color)
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_colored_bar(document, text, fill='2563eb', color='ffffff'):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_text(cell, text, bold=True, color=color, size=12)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    document.add_paragraph()

def add_heading(document, text, level=1, color='0f172a'):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = rgb(color)
    run.font.size = Pt(22 if level == 1 else 15)
    return p

def add_body(document, text, bold=False, color='334155', size=10.5):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = rgb(color)
    run.font.size = Pt(size)
    return p

def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        run.font.color.rgb = rgb('334155')

def add_card_grid(document, cards, cols=2):
    rows = (len(cards) + cols - 1) // cols
    table = document.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    idx = 0
    for r in range(rows):
        for c in range(cols):
            cell = table.cell(r, c)
            shade_cell(cell, 'eef6ff' if (r + c) % 2 == 0 else 'f0fdf4')
            if idx < len(cards):
                title, body, accent = cards[idx]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(title)
                run.bold = True
                run.font.color.rgb = rgb(accent)
                run.font.size = Pt(11)
                p2 = cell.add_paragraph()
                run2 = p2.add_run(body)
                run2.font.color.rgb = rgb('334155')
                run2.font.size = Pt(9.5)
            else:
                cell.text = ''
            idx += 1
    document.add_paragraph()

def create_architecture_image(path):
    w, h = 1600, 900
    img = Image.new('RGB', (w, h), '#07111f')
    d = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype('arialbd.ttf', 54)
        head_font = ImageFont.truetype('arialbd.ttf', 34)
        body_font = ImageFont.truetype('arial.ttf', 25)
    except Exception:
        title_font = head_font = body_font = ImageFont.load_default()

    for i in range(0, w, 40):
        color = (12, 27, 48) if (i // 40) % 2 == 0 else (8, 20, 36)
        d.line((i, 0, i + 260, h), fill=color, width=2)

    d.text((70, 55), 'NeuralBox Runtime Map', fill='#ffffff', font=title_font)
    d.text((72, 122), 'Browser-only AI: model, memory, docs, voice, and trust live client-side.', fill='#93c5fd', font=body_font)

    boxes = [
        ('HTML/CSS Shell', 90, 240, '#38bdf8'),
        ('main.js Orchestrator', 590, 240, '#60a5fa'),
        ('WebLLM + WebGPU', 1090, 240, '#34d399'),
        ('IndexedDB Storage', 90, 545, '#fbbf24'),
        ('RAG + Web Search', 590, 545, '#a78bfa'),
        ('Voice + Vision', 1090, 545, '#fb7185'),
    ]
    for label, x, y, color in boxes:
        d.rounded_rectangle((x, y, x + 380, y + 150), radius=28, fill='#0f1c30', outline=color, width=5)
        d.text((x + 28, y + 32), label, fill=color, font=head_font)
        sub = {
            'HTML/CSS Shell': 'DOM ids stay stable\nresponsive UI layer',
            'main.js Orchestrator': 'state, events, flows\nstreaming generation',
            'WebLLM + WebGPU': 'local model runtime\nprivate inference',
            'IndexedDB Storage': 'settings, chats, docs\nlegacy migration',
            'RAG + Web Search': 'local context + optional\nnetwork enrichment',
            'Voice + Vision': 'Whisper, TTS, images\ncapability gated',
        }[label]
        d.multiline_text((x + 28, y + 82), sub, fill='#cbd5e1', font=body_font, spacing=6)

    arrows = [((470, 315), (590, 315)), ((970, 315), (1090, 315)), ((280, 545), (280, 390)), ((780, 545), (780, 390)), ((1280, 545), (1280, 390))]
    for start, end in arrows:
        d.line((*start, *end), fill='#e2e8f0', width=5)
        ex, ey = end
        d.polygon([(ex, ey), (ex - 18, ey - 10), (ex - 18, ey + 10)], fill='#e2e8f0')

    d.rounded_rectangle((390, 760, 1210, 835), radius=22, fill='#e0f2fe', outline='#38bdf8', width=3)
    d.text((430, 780), 'Trust layer: route reason, model metadata, web/RAG citations, confidence.', fill='#075985', font=body_font)
    img.save(path)

def create_validation_image(path):
    w, h = 1600, 720
    img = Image.new('RGB', (w, h), '#f8fafc')
    d = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype('arialbd.ttf', 50)
        badge_font = ImageFont.truetype('arialbd.ttf', 25)
        body_font = ImageFont.truetype('arial.ttf', 22)
    except Exception:
        title_font = badge_font = body_font = ImageFont.load_default()
    d.text((70, 55), 'Validation Wall of Green', fill='#0f172a', font=title_font)
    d.text((72, 118), 'Static contracts, helper tests, RAG web ingest, production build, and browser lifecycle all pass.', fill='#475569', font=body_font)
    badges = [
        'audit', 'env', 'models', 'composer', 'generation', 'events', 'voice', 'settings',
        'RAG helpers', 'web search', 'accessibility', 'rendering', 'routing', 'device',
        'trust', 'ASCII UI', 'stability', 'RAG web', 'build', 'browser smoke'
    ]
    x, y = 80, 210
    for i, badge in enumerate(badges):
        d.rounded_rectangle((x, y, x + 260, y + 72), radius=22, fill='#dcfce7', outline='#22c55e', width=3)
        d.text((x + 26, y + 22), f'PASS: {badge}', fill='#166534', font=badge_font)
        x += 300
        if x + 260 > w - 70:
            x = 80
            y += 102
    d.rounded_rectangle((80, 625, 1520, 675), radius=18, fill='#0f172a')
    d.text((110, 638), 'Honest limit: real WebGPU inference, microphone, TTS, and vision still deserve manual device testing.', fill='#e2e8f0', font=body_font)
    img.save(path)

arch_path = ASSET_DIR / 'architecture.png'
validation_path = ASSET_DIR / 'validation.png'
create_architecture_image(arch_path)
create_validation_image(validation_path)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)

styles = doc.styles
styles['Normal'].font.name = 'Aptos'
styles['Normal'].font.size = Pt(10.5)

# Cover
add_colored_bar(doc, 'BUILDER CLUB TECH SHOWCASE', '07111f')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('NeuralBox')
r.bold = True
r.font.size = Pt(44)
r.font.color.rgb = rgb('2563eb')
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run('Private AI in the Browser')
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = rgb('0f172a')
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p3.add_run('No server. No account. WebGPU-powered local chat with RAG, voice, vision, and trust metadata.')
r.font.size = Pt(12)
r.font.color.rgb = rgb('475569')
add_card_grid(doc, [
    ('Core idea', 'Run useful AI locally, inside a browser tab.', '2563eb'),
    ('Builder energy', 'A serious app with hackable surfaces everywhere.', '16a34a'),
    ('Technical angle', 'WebGPU, WebLLM, IndexedDB, RAG, Whisper, Playwright.', '7c3aed'),
    ('Honest status', 'Much stronger now, still with real engineering frontiers.', 'd97706'),
], cols=2)

# Pitch
add_heading(doc, '1. The 10-second pitch')
add_body(doc, 'NeuralBox is a local-first AI workbench that runs models directly in the browser using WebGPU. It keeps chats and documents local, adds optional web lookup only when requested, and explains why each answer happened through trust metadata.')
add_bullets(doc, [
    'Private by design: conversations and RAG documents stay in browser storage.',
    'No backend app server: Vite serves a single-page app; inference happens client-side.',
    'Expandable interface: text, local docs, optional web search, voice, and image input.',
    'Builder-friendly: plain JavaScript, modular helpers, smoke tests, and clear docs.',
])

# Architecture
add_heading(doc, '2. Architecture: tiny app shell, surprisingly big runtime')
doc.add_picture(str(arch_path), width=Inches(7.2))
add_body(doc, 'The big technical trick is not just local inference. It is coordinating model lifecycle, GPU fit, local persistence, optional retrieval, dynamic UI state, and diagnostics in one browser app.')

# Feature tour
add_heading(doc, '3. Feature tour: what builders can poke at')
add_card_grid(doc, [
    ('Local chat', 'Streaming responses from WebLLM through WebGPU.', '2563eb'),
    ('Auto routing', 'Task heuristics can hot-swap to a better model.', '0891b2'),
    ('Local RAG', 'Docs are chunked, scored, cited, and kept local.', '16a34a'),
    ('RAG profiles', 'Precise, Balanced, and Broad retrieval modes.', '15803d'),
    ('Web recovery', 'Search failures classify and degrade gracefully.', 'd97706'),
    ('Voice', 'Whisper transcription plus browser speech synthesis.', 'be123c'),
    ('Vision', 'Image input for supported models with compatibility handling.', '7c3aed'),
    ('Trust layer', 'Model, route, sources, profile, and confidence metadata.', '0f766e'),
], cols=2)

# Improvements
add_heading(doc, '4. What changed in the improvement sweep')
rows = [
    ('Security', 'Audit cleaned to 0 vulnerabilities; Vite lockfile moved to safe resolved version.'),
    ('Accessibility', 'Dialogs, live regions, labels, aria state, keyboard activation, and focus styling.'),
    ('Web search', 'Failure classification, recovery notices, and workbench/runtime telemetry.'),
    ('RAG', 'Configurable retrieval profiles plus trust/export metadata.'),
    ('Architecture', 'Model catalog extracted into `src/lib/models.js` with integrity test.'),
    ('Documentation', 'Codebase scan, improvement report, test report, roadmap/task sync.'),
]
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
shade_cell(t.cell(0,0), 'dbeafe')
shade_cell(t.cell(0,1), 'dbeafe')
set_cell_text(t.cell(0,0), 'Area', True, '1e3a8a', 10)
set_cell_text(t.cell(0,1), 'Upgrade', True, '1e3a8a', 10)
for area, upgrade in rows:
    row = t.add_row()
    set_cell_text(row.cells[0], area, True, '0f172a', 9.5)
    row.cells[1].text = upgrade
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
    row.cells[1].paragraphs[0].runs[0].font.color.rgb = rgb('334155')
doc.add_paragraph()

# Validation
add_heading(doc, '5. Testing: the validation wall')
doc.add_picture(str(validation_path), width=Inches(7.2))
add_body(doc, 'Automated tests validate helper logic, static UI contracts, source safety, RAG ingestion, production build, and a Playwright browser lifecycle. The suite is not a replacement for real GPU demo testing, but it is a strong regression net.')

# Demo script
add_heading(doc, '6. Builder club demo script')
add_bullets(doc, [
    'Open NeuralBox and explain: this is not a hosted chatbot; it is browser-local inference.',
    'Show model selector: tiny models for low memory, stronger models for better GPUs, Auto mode for routing.',
    'Ask a normal chat question and point out streaming plus local storage.',
    'Attach RAG docs and switch between Precise, Balanced, and Broad retrieval.',
    'Open trust metadata under an answer: model, route, RAG docs, confidence, web mode.',
    'Toggle web search and intentionally discuss graceful failure if the endpoint is unavailable.',
    'Show voice input or voice overlay if microphone permissions cooperate.',
    'Close with the builder challenge: what would you extract, improve, or ship next?',
])

# Honest engineering
add_heading(doc, '7. The honest engineering story')
add_body(doc, 'The codebase is now stronger, but the most interesting part is that the hard problems are visible instead of hidden.')
add_card_grid(doc, [
    ('Still hard', 'Real WebGPU behavior depends on device, browser, VRAM, and upstream WebLLM.', 'dc2626'),
    ('Still big', '`src/main.js` remains the central orchestrator and should be split further.', 'd97706'),
    ('Still lexical', 'RAG uses token scoring, not semantic embeddings yet.', '7c3aed'),
    ('Still fun', 'Every remaining gap is a great builder-club challenge.', '16a34a'),
], cols=2)

# Builder challenges
add_heading(doc, '8. Next builder challenges')
add_bullets(doc, [
    'Extract feature controllers: settings, conversations, web search runtime, RAG UI, voice chat.',
    'Add embedding-based local retrieval and compare it against lexical RAG profiles.',
    'Create a GPU benchmark screen that recommends models from measured tokens/sec.',
    'Add a first-run demo dataset and guided onboarding tour.',
    'Build a local plugin system for tools that still preserve privacy boundaries.',
])

# Appendix
add_heading(doc, 'Appendix: repo guide for technical readers')
add_body(doc, 'Important files to inspect:', bold=True, color='0f172a')
add_bullets(doc, [
    'Runtime orchestrator: `src/main.js`',
    'Model catalog: `src/lib/models.js`',
    'RAG helpers/profiles: `src/lib/rag.js`',
    'Web recovery helpers: `src/lib/web-search.js`',
    'Persistence: `src/db/database.js`',
    'Validation report: `doc/TEST_REPORT_2026-05-04.md`',
    'Improvement report: `doc/IMPROVEMENT_REPORT_2026-05-04.md`',
])
add_body(doc, 'Suggested closing line: NeuralBox is not just a chatbot demo. It is a browser-native AI systems lab wearing a friendly chat UI.')

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
